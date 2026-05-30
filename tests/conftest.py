"""Общие фикстуры и хелперы для тестов."""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document as DocxDocument


def make_docx(blocks: list[tuple[int, str]]) -> bytes:
    """Собирает .docx в памяти из списка (level, text).

    level=0 — обычный абзац, level>=1 — заголовок соответствующего уровня.
    Возвращает байты файла.
    """
    doc = DocxDocument()
    for level, text in blocks:
        if level >= 1:
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_docx_with_table(heading: str, rows: list[list[str]]) -> bytes:
    """Собирает .docx с одним заголовком и таблицей."""
    doc = DocxDocument()
    doc.add_heading(heading, level=1)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.cell(r, c).text = val
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def contract_a() -> bytes:
    return make_docx([
        (1, "Предмет договора"),
        (0, "Исполнитель оказывает консультационные услуги Заказчику."),
        (1, "Стоимость и оплата"),
        (0, "Стоимость услуг составляет 100000 рублей. Срок оплаты — 10 дней."),
        (1, "Срок действия"),
        (0, "Договор действует до 31 декабря 2026 года."),
    ])


@pytest.fixture
def contract_b() -> bytes:
    return make_docx([
        (1, "Предмет договора"),
        (0, "Исполнитель оказывает консультационные услуги Заказчику."),
        (1, "Стоимость и оплата"),
        (0, "Стоимость услуг составляет 150000 рублей. Срок оплаты — 5 дней."),
        (1, "Срок действия"),
        (0, "Договор действует до 31 декабря 2026 года."),
        (1, "Ответственность сторон"),
        (0, "За просрочку оплаты начисляется штраф 1% за каждый день."),
    ])
